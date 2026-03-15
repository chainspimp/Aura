# ============================================
# FILE: tools/osint.py
# Deep OSINT Engine v2 — concurrent, multi-source
# All public data only. No auth bypass.
# ============================================

import re
import time
import hashlib
import logging
import requests
import concurrent.futures
from datetime import datetime
from typing import Optional, List, Dict, Tuple
from urllib.parse import quote_plus

logger = logging.getLogger(__name__)

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
}
TIMEOUT = 12


# ═══════════════════════════════════════════════════════════
# USERNAME PERMUTATION ENGINE
# Generates likely usernames from a real name
# ═══════════════════════════════════════════════════════════

def generate_username_variants(name: str) -> List[str]:
    """Generate common username patterns from a full name."""
    if not name:
        return []
    parts = name.lower().split()
    if len(parts) < 2:
        return [parts[0]]

    first, last = parts[0], parts[-1]
    mid = parts[1] if len(parts) > 2 else ""

    variants = [
        first + last,
        first + "." + last,
        first + "_" + last,
        first[0] + last,
        first[0] + "." + last,
        first[0] + "_" + last,
        first + last[0],
        first + "." + last[0],
        last + first,
        last + "." + first,
        last + "_" + first,
        last + first[0],
        first,
        last,
        first + last[:3],
        first[:3] + last,
    ]
    if mid:
        variants += [
            first + mid[0] + last,
            first[0] + mid[0] + last,
        ]

    # Deduplicate preserving order
    seen = set()
    out = []
    for v in variants:
        if v not in seen and len(v) >= 3:
            seen.add(v)
            out.append(v)
    return out[:20]


# ═══════════════════════════════════════════════════════════
# PLATFORM USERNAME CHECKER
# Checks 40+ platforms for a username via HTTP probing
# ═══════════════════════════════════════════════════════════

# Format: (platform_name, url_template, check_type, false_positive_strings)
# check_type:
#   status_200         — exists if HTTP 200
#   not_in_body:TEXT   — exists if 200 AND "TEXT" NOT in response body
#   in_body:TEXT       — exists if 200 AND "TEXT" IS in response body
#   api_json:KEY       — exists if 200 AND json response has truthy KEY
#
# false_positive_strings: page content that means "user not found" even on 200
PLATFORMS = [
    # Social — most return 200 for missing users, need body checks
    ("Twitter/X",    "https://twitter.com/{}",                    "not_in_body:This account doesn\'t exist"),
    ("Instagram",    "https://www.instagram.com/{}/",             "not_in_body:Sorry, this page isn\'t available"),
    ("TikTok",       "https://www.tiktok.com/@{}",                "not_in_body:Couldn\'t find this account"),
    ("Facebook",     "https://www.facebook.com/{}",               "not_in_body:The link you followed may be broken"),
    ("Pinterest",    "https://www.pinterest.com/{}/",             "not_in_body:Sorry, we couldn\'t find that page"),
    ("Tumblr",       "https://{}.tumblr.com/",                    "not_in_body:There\'s nothing here."),
    ("Snapchat",     "https://www.snapchat.com/add/{}",           "not_in_body:Sorry, we couldn\'t find"),
    ("Mastodon",     "https://mastodon.social/@{}",               "not_in_body:The page you are looking for"),
    ("Bluesky",      "https://bsky.app/profile/{}",               "not_in_body:Profile not found"),
    ("Threads",      "https://www.threads.net/@{}",               "not_in_body:Page not found"),
    # Dev — APIs are more reliable
    ("GitHub",       "https://api.github.com/users/{}",           "api_json:login"),
    ("GitLab",       "https://gitlab.com/api/v4/users?username={}","in_body:\"username\""),
    ("Bitbucket",    "https://bitbucket.org/{}",                  "not_in_body:No Bitbucket account"),
    ("HackerNews",   "https://hacker-news.firebaseio.com/v0/user/{}.json", "not_in_body:null"),
    ("CodePen",      "https://codepen.io/{}",                     "not_in_body:Sorry, we couldn\'t find"),
    ("Replit",       "https://replit.com/@{}",                    "not_in_body:User not found"),
    ("npm",          "https://registry.npmjs.org/-/v1/search?text=author:{}&size=1", "in_body:\"objects\""),
    ("PyPI",         "https://pypi.org/user/{}/",                 "not_in_body:404 Not Found"),
    ("Docker Hub",   "https://hub.docker.com/v2/users/{}/",       "api_json:username"),
    # Gaming — use body checks to avoid false positives
    ("Steam",        "https://steamcommunity.com/id/{}",          "not_in_body:The specified profile could not be found"),
    ("Twitch",       "https://www.twitch.tv/{}",                  "not_in_body:Sorry. Unless you\'ve got a time machine"),
    ("Roblox",       "https://users.roblox.com/v1/users/search?keyword={}&limit=10", "in_body:\"data\""),
    ("Chess.com",    "https://api.chess.com/pub/player/{}",       "api_json:username"),
    # Creative
    ("Medium",       "https://medium.com/@{}",                    "not_in_body:Page not found"),
    ("Substack",     "https://substack.com/@{}",                  "not_in_body:This publication does not exist"),
    ("Dev.to",       "https://dev.to/api/users/by_username?url={}", "api_json:id"),
    ("Behance",      "https://www.behance.net/{}",                "not_in_body:The page you requested could not be found"),
    ("Dribbble",     "https://dribbble.com/{}",                   "not_in_body:Whoops, that page is gone"),
    ("SoundCloud",   "https://soundcloud.com/{}",                 "not_in_body:We can\'t find that user"),
    ("Bandcamp",     "https://{}.bandcamp.com/",                  "not_in_body:Sorry, that something isn\'t here"),
    ("Flickr",       "https://www.flickr.com/people/{}",          "not_in_body:This member does not exist"),
    ("Vimeo",        "https://vimeo.com/{}",                      "not_in_body:Sorry, we couldn\'t find that page"),
    ("YouTube",      "https://www.youtube.com/@{}",               "not_in_body:This page isn\'t available"),
    # Professional
    ("LinkedIn",     "https://www.linkedin.com/in/{}",            "not_in_body:Page not found"),
    ("Keybase",      "https://keybase.io/_/api/1.0/user/lookup.json?usernames={}", "not_in_body:\"them\":[null]"),
    ("Product Hunt", "https://www.producthunt.com/@{}",           "not_in_body:Oops! This page doesn\'t exist"),
    ("About.me",     "https://about.me/{}",                       "not_in_body:Sorry, this page doesn\'t exist"),
    # Other
    ("Pastebin",     "https://pastebin.com/u/{}",                 "not_in_body:Not Found"),
    ("Reddit",       "https://www.reddit.com/user/{}/about.json", "not_in_body:\"error\""),
]


def check_username_on_platform(platform: str, url_template: str,
                                 check_type: str, username: str) -> Optional[Dict]:
    """
    Check if a username exists on a single platform using smart detection.
    Returns a result dict if found, None otherwise.
    """
    url = url_template.format(username)
    # Use the original URL as the display URL (not the API endpoint)
    display_url = url

    try:
        r = requests.get(url, headers=HEADERS, timeout=8, allow_redirects=True)
        exists = False

        if check_type == "status_200":
            exists = r.status_code == 200

        elif check_type.startswith("not_in_body:"):
            needle = check_type.split(":", 1)[1]
            exists = r.status_code == 200 and needle.lower() not in r.text.lower()

        elif check_type.startswith("in_body:"):
            needle = check_type.split(":", 1)[1]
            exists = r.status_code == 200 and needle.lower() in r.text.lower()

        elif check_type.startswith("api_json:"):
            key = check_type.split(":", 1)[1]
            if r.status_code == 200:
                try:
                    data = r.json()
                    # Handle both dict and list responses
                    if isinstance(data, list):
                        exists = len(data) > 0 and data[0] is not None and bool(data[0].get(key) if isinstance(data[0], dict) else True)
                    else:
                        exists = bool(data.get(key))
                except Exception:
                    exists = False

        if exists:
            # Normalise display URL to profile page (not API endpoint)
            profile_url = url
            api_prefixes = [
                "https://api.github.com/users/",
                "https://gitlab.com/api/v4/users?username=",
                "https://hacker-news.firebaseio.com/v0/user/",
                "https://registry.npmjs.org/-/v1/search",
                "https://hub.docker.com/v2/users/",
                "https://users.roblox.com/v1/users/search",
                "https://api.chess.com/pub/player/",
                "https://dev.to/api/users/by_username",
                "https://www.reddit.com/user/",
                "https://keybase.io/_/api/",
            ]
            profile_map = {
                "GitHub":     f"https://github.com/{username}",
                "GitLab":     f"https://gitlab.com/{username}",
                "HackerNews": f"https://news.ycombinator.com/user?id={username}",
                "npm":        f"https://www.npmjs.com/~{username}",
                "Docker Hub": f"https://hub.docker.com/u/{username}",
                "Roblox":     f"https://www.roblox.com/users/profile?username={username}",
                "Chess.com":  f"https://www.chess.com/member/{username}",
                "Dev.to":     f"https://dev.to/{username}",
                "Reddit":     f"https://www.reddit.com/user/{username}",
                "Keybase":    f"https://keybase.io/{username}",
            }
            profile_url = profile_map.get(platform, url)
            return {"platform": platform, "url": profile_url, "username": username}

    except Exception:
        pass
    return None


def scan_username_across_platforms(username: str,
                                    progress_cb=None) -> List[Dict]:
    """
    Concurrently check a username across all platforms.
    Returns list of found profiles.
    """
    found = []
    total = len(PLATFORMS)

    with concurrent.futures.ThreadPoolExecutor(max_workers=20) as ex:
        futures = {
            ex.submit(check_username_on_platform, p, u, ct, username): p
            for p, u, ct in PLATFORMS
        }
        done = 0
        for future in concurrent.futures.as_completed(futures):
            done += 1
            result = future.result()
            if result:
                found.append(result)
                if progress_cb:
                    progress_cb(f"   ✅ Found on {result['platform']}: {result['url']}")
            if progress_cb and done % 10 == 0:
                progress_cb(f"   Checked {done}/{total} platforms...")

    return sorted(found, key=lambda x: x["platform"])


# ═══════════════════════════════════════════════════════════
# MAIN OSINT ENGINE
# ═══════════════════════════════════════════════════════════

class OSINTEngine:
    """
    Deep public OSINT aggregator.
    Concurrent platform scanning, username permutation,
    rich API lookups, and targeted web searches.
    """

    def __init__(self, web_search_fn=None):
        self._search   = web_search_fn
        self.results   = {}
        self.errors    = []
        self._progress = print   # can be replaced with GUI callback

    def set_progress_callback(self, cb):
        self._progress = cb

    def _log(self, msg):
        self._progress(msg)

    # ── Main entry point ──────────────────────────────────────────────────────

    def investigate(self,
                    name:     Optional[str] = None,
                    email:    Optional[str] = None,
                    username: Optional[str] = None,
                    location: Optional[str] = None,
                    phone:    Optional[str] = None,
                    employer: Optional[str] = None,
                    age:      Optional[str] = None,
                    website:  Optional[str] = None) -> dict:

        if not any([name, email, username]):
            return {"error": "Provide at least one of: name, email, username"}

        self.results = {
            "query": {k: v for k, v in {
                "name": name, "email": email, "username": username,
                "location": location, "phone": phone,
                "employer": employer, "age": age, "website": website,
                "timestamp": datetime.now().isoformat(),
            }.items() if v},
            "platform_hits":   [],   # from username scanner
            "api_profiles":    {},   # from rich API lookups
            "social_profiles": {},
            "data_breaches":   {},
            "web_presence":    {},
            "phone_info":      {},
            "domain_info":     {},
            "image_search":    {},
            "summary":         [],
        }
        self.errors = []

        # ── 1. Username platform sweep (concurrent) ──
        usernames_to_check = []
        if username:
            usernames_to_check.append(username)
        if name and not username:
            # Generate variants and check top ones
            variants = generate_username_variants(name)
            self._log(f"   Generated {len(variants)} username variants from name")
            usernames_to_check.extend(variants[:8])

        if usernames_to_check:
            self._log(f"🌐 Scanning {len(PLATFORMS)} platforms for {len(usernames_to_check)} username(s)...")
            all_hits = []
            for uname in usernames_to_check:
                self._log(f"   Checking username: {uname}")
                hits = scan_username_across_platforms(uname, self._log)
                all_hits.extend(hits)
            # Deduplicate by platform
            seen_platforms = set()
            unique_hits = []
            for h in all_hits:
                if h["platform"] not in seen_platforms:
                    seen_platforms.add(h["platform"])
                    unique_hits.append(h)
            self.results["platform_hits"] = unique_hits
            self._log(f"   ✅ Found on {len(unique_hits)} platforms")

        # ── 2. Rich API lookups ──
        primary_username = username or (generate_username_variants(name)[0] if name else None)

        if primary_username:
            self._api_github(primary_username)
            self._api_reddit(primary_username)
            self._api_gitlab(primary_username)
            self._api_keybase(primary_username)
            self._api_npm(primary_username)
            self._api_stackoverflow(primary_username)
            self._api_hackernews(primary_username)
            self._api_devto(primary_username)

        # ── 3. Email lookups ──
        if email:
            self._check_breaches(email)
            self._check_gravatar(email)
            self._check_email_format(email)

        # ── 4. Phone lookup ──
        if phone:
            self._check_phone(phone)

        # ── 5. Website/domain lookup ──
        if website:
            self._check_domain(website)

        # ── 6. Deep web searches ──
        self._deep_web_search(name, username, email, location, employer, phone, website)

        # ── 7. Image/reverse search links ──
        if name:
            self._generate_image_search_links(name, location)

        self._build_summary()
        return self.results

    # ── Rich API checks ───────────────────────────────────────────────────────

    def _api_github(self, username: str):
        self._log(f"   📦 GitHub API: {username}")
        try:
            r = requests.get(f"https://api.github.com/users/{username}",
                             headers=HEADERS, timeout=TIMEOUT)
            if r.status_code != 200:
                return
            d = r.json()

            # Get repos
            repos = []
            rr = requests.get(
                f"https://api.github.com/users/{username}/repos?per_page=100&sort=pushed",
                headers=HEADERS, timeout=TIMEOUT)
            if rr.status_code == 200:
                for repo in rr.json():
                    repos.append({
                        "name":     repo["name"],
                        "desc":     repo.get("description", ""),
                        "lang":     repo.get("language"),
                        "stars":    repo.get("stargazers_count", 0),
                        "forks":    repo.get("forks_count", 0),
                        "url":      repo.get("html_url"),
                        "updated":  repo.get("pushed_at", "")[:10],
                        "topics":   repo.get("topics", []),
                    })

            # Get starred repos (reveals interests)
            starred = []
            rs = requests.get(
                f"https://api.github.com/users/{username}/starred?per_page=10",
                headers=HEADERS, timeout=TIMEOUT)
            if rs.status_code == 200:
                for s in rs.json()[:5]:
                    starred.append(s.get("full_name", ""))

            # Get orgs
            orgs = []
            ro = requests.get(
                f"https://api.github.com/users/{username}/orgs",
                headers=HEADERS, timeout=TIMEOUT)
            if ro.status_code == 200:
                orgs = [o.get("login") for o in ro.json()]

            # Get events (recent activity)
            activity = []
            re_ = requests.get(
                f"https://api.github.com/users/{username}/events/public?per_page=10",
                headers=HEADERS, timeout=TIMEOUT)
            if re_.status_code == 200:
                for ev in re_.json()[:5]:
                    activity.append({
                        "type": ev.get("type"),
                        "repo": ev.get("repo", {}).get("name"),
                        "date": ev.get("created_at", "")[:10],
                    })

            total_stars = sum(r["stars"] for r in repos)
            langs = {}
            for repo in repos:
                if repo["lang"]:
                    langs[repo["lang"]] = langs.get(repo["lang"], 0) + 1
            top_langs = sorted(langs.items(), key=lambda x: x[1], reverse=True)[:5]

            self.results["api_profiles"]["github"] = {
                "found":         True,
                "url":           d.get("html_url"),
                "name":          d.get("name"),
                "bio":           d.get("bio"),
                "location":      d.get("location"),
                "company":       d.get("company"),
                "blog":          d.get("blog"),
                "email":         d.get("email"),
                "twitter":       d.get("twitter_username"),
                "followers":     d.get("followers"),
                "following":     d.get("following"),
                "public_repos":  d.get("public_repos"),
                "public_gists":  d.get("public_gists"),
                "created_at":    d.get("created_at", "")[:10],
                "updated_at":    d.get("updated_at", "")[:10],
                "hireable":      d.get("hireable"),
                "total_stars":   total_stars,
                "top_languages": [l[0] for l in top_langs],
                "organizations": orgs,
                "starred_repos": starred,
                "top_repos":     sorted(repos, key=lambda x: x["stars"], reverse=True)[:10],
                "recent_activity": activity,
            }
            self._log(f"      ✅ {d.get('name', username)} — {len(repos)} repos, {total_stars} total stars")
        except Exception as e:
            self.errors.append(f"GitHub API: {e}")

    def _api_reddit(self, username: str):
        self._log(f"   🟠 Reddit API: {username}")
        try:
            r = requests.get(
                f"https://www.reddit.com/user/{username}/about.json",
                headers={**HEADERS, "User-Agent": "AURA-OSINT/2.0"},
                timeout=TIMEOUT)
            if r.status_code != 200:
                return
            d = r.json().get("data", {})

            # Also get recent posts
            posts = []
            rp = requests.get(
                f"https://www.reddit.com/user/{username}/submitted.json?limit=10",
                headers={**HEADERS, "User-Agent": "AURA-OSINT/2.0"},
                timeout=TIMEOUT)
            if rp.status_code == 200:
                for post in rp.json().get("data", {}).get("children", [])[:5]:
                    pd = post.get("data", {})
                    posts.append({
                        "title":     pd.get("title", ""),
                        "subreddit": pd.get("subreddit", ""),
                        "score":     pd.get("score", 0),
                        "url":       f"https://reddit.com{pd.get('permalink', '')}",
                        "date":      datetime.utcfromtimestamp(
                                         pd.get("created_utc", 0)).strftime("%Y-%m-%d"),
                    })

            # Get top subreddits from comments
            subs = []
            rc = requests.get(
                f"https://www.reddit.com/user/{username}/comments.json?limit=100",
                headers={**HEADERS, "User-Agent": "AURA-OSINT/2.0"},
                timeout=TIMEOUT)
            if rc.status_code == 200:
                sub_counts = {}
                for c in rc.json().get("data", {}).get("children", []):
                    sub = c.get("data", {}).get("subreddit", "")
                    if sub:
                        sub_counts[sub] = sub_counts.get(sub, 0) + 1
                subs = sorted(sub_counts.items(), key=lambda x: x[1], reverse=True)[:10]

            self.results["api_profiles"]["reddit"] = {
                "found":            True,
                "url":              f"https://reddit.com/u/{username}",
                "karma_post":       d.get("link_karma", 0),
                "karma_comment":    d.get("comment_karma", 0),
                "total_karma":      d.get("total_karma", 0),
                "created":          datetime.utcfromtimestamp(
                                        d.get("created_utc", 0)
                                    ).strftime("%Y-%m-%d") if d.get("created_utc") else None,
                "is_mod":           d.get("is_mod", False),
                "verified":         d.get("verified", False),
                "has_premium":      d.get("is_gold", False),
                "top_subreddits":   [f"r/{s[0]} ({s[1]})" for s in subs],
                "recent_posts":     posts,
            }
            self._log(f"      ✅ {d.get('total_karma', 0)} karma, active in {len(subs)} subreddits")
        except Exception as e:
            self.errors.append(f"Reddit API: {e}")

    def _api_hackernews(self, username: str):
        self._log(f"   🔶 HackerNews API: {username}")
        try:
            r = requests.get(
                f"https://hacker-news.firebaseio.com/v0/user/{username}.json",
                headers=HEADERS, timeout=TIMEOUT)
            if r.status_code != 200 or not r.json():
                return
            d = r.json()
            self.results["api_profiles"]["hackernews"] = {
                "found":    True,
                "url":      f"https://news.ycombinator.com/user?id={username}",
                "karma":    d.get("karma", 0),
                "about":    re.sub(r'<[^>]+>', '', d.get("about", "")),
                "created":  datetime.utcfromtimestamp(
                                d.get("created", 0)).strftime("%Y-%m-%d"),
                "submitted_count": len(d.get("submitted", [])),
            }
            self._log(f"      ✅ {d.get('karma', 0)} karma")
        except Exception as e:
            self.errors.append(f"HackerNews: {e}")

    def _api_devto(self, username: str):
        self._log(f"   💻 Dev.to API: {username}")
        try:
            r = requests.get(
                f"https://dev.to/api/users/by_username?url={username}",
                headers=HEADERS, timeout=TIMEOUT)
            if r.status_code != 200:
                return
            d = r.json()
            if not d.get("id"):
                return

            # Get articles
            articles = []
            ra = requests.get(
                f"https://dev.to/api/articles?username={username}&per_page=5",
                headers=HEADERS, timeout=TIMEOUT)
            if ra.status_code == 200:
                for a in ra.json()[:5]:
                    articles.append({
                        "title":        a.get("title"),
                        "reactions":    a.get("positive_reactions_count", 0),
                        "comments":     a.get("comments_count", 0),
                        "published":    a.get("published_at", "")[:10],
                        "url":          a.get("url"),
                        "tags":         a.get("tag_list", []),
                    })

            self.results["api_profiles"]["devto"] = {
                "found":           True,
                "url":             f"https://dev.to/{username}",
                "name":            d.get("name"),
                "summary":         d.get("summary"),
                "location":        d.get("location"),
                "github":          d.get("github_username"),
                "twitter":         d.get("twitter_username"),
                "website":         d.get("website_url"),
                "joined":          d.get("joined_at", "")[:10],
                "articles_count":  len(articles),
                "top_articles":    articles,
            }
            self._log(f"      ✅ {d.get('name')} — {len(articles)} articles")
        except Exception as e:
            self.errors.append(f"Dev.to: {e}")

    def _api_gitlab(self, username: str):
        self._log(f"   🦊 GitLab API: {username}")
        try:
            r = requests.get(
                f"https://gitlab.com/api/v4/users?username={username}",
                headers=HEADERS, timeout=TIMEOUT)
            if r.status_code != 200 or not r.json():
                return
            u = r.json()[0]
            uid = u.get("id")

            # Get projects
            projects = []
            rp = requests.get(
                f"https://gitlab.com/api/v4/users/{uid}/projects?per_page=10",
                headers=HEADERS, timeout=TIMEOUT)
            if rp.status_code == 200:
                for p in rp.json()[:5]:
                    projects.append({
                        "name":  p.get("name"),
                        "desc":  p.get("description", ""),
                        "stars": p.get("star_count", 0),
                        "url":   p.get("web_url"),
                        "lang":  p.get("predominant_language"),
                    })

            self.results["api_profiles"]["gitlab"] = {
                "found":    True,
                "url":      u.get("web_url"),
                "name":     u.get("name"),
                "bio":      u.get("bio"),
                "location": u.get("location"),
                "website":  u.get("website_url"),
                "created":  u.get("created_at", "")[:10],
                "projects": projects,
            }
            self._log(f"      ✅ {u.get('name')}")
        except Exception as e:
            self.errors.append(f"GitLab: {e}")

    def _api_keybase(self, username: str):
        self._log(f"   🔑 Keybase: {username}")
        try:
            r = requests.get(
                f"https://keybase.io/_/api/1.0/user/lookup.json?usernames={username}",
                headers=HEADERS, timeout=TIMEOUT)
            if r.status_code != 200:
                return
            them = r.json().get("them", [])
            # API returns [null] for non-existent users
            if not them or them[0] is None:
                return
            u = them[0]
            if not isinstance(u, dict):
                return
            proofs = u.get("proofs_summary", {}).get("all", [])

            self.results["api_profiles"]["keybase"] = {
                "found":            True,
                "url":              f"https://keybase.io/{username}",
                "full_name":        u.get("profile", {}).get("full_name"),
                "bio":              u.get("profile", {}).get("bio"),
                "location":         u.get("profile", {}).get("location"),
                "verified_proofs":  [
                    {"service": p.get("proof_type"),
                     "username": p.get("nametag"),
                     "url": p.get("service_url")}
                    for p in proofs
                ],
                "has_pgp_key":      bool(u.get("public_keys", {}).get("primary")),
            }
            self._log(f"      ✅ {len(proofs)} verified identities linked")
        except Exception as e:
            self.errors.append(f"Keybase: {e}")

    def _api_npm(self, username: str):
        self._log(f"   📦 npm: {username}")
        try:
            r = requests.get(
                f"https://registry.npmjs.org/-/v1/search?text=author:{username}&size=20",
                headers=HEADERS, timeout=TIMEOUT)
            if r.status_code != 200:
                return
            pkgs = r.json().get("objects", [])
            if not pkgs:
                return
            total_dl = 0
            package_list = []
            for p in pkgs:
                pkg = p["package"]
                package_list.append({
                    "name":    pkg["name"],
                    "desc":    pkg.get("description", ""),
                    "version": pkg.get("version"),
                    "date":    pkg.get("date", "")[:10],
                    "links":   pkg.get("links", {}),
                })
            self.results["api_profiles"]["npm"] = {
                "found":    True,
                "url":      f"https://www.npmjs.com/~{username}",
                "packages": package_list,
                "count":    len(package_list),
            }
            self._log(f"      ✅ {len(package_list)} packages")
        except Exception as e:
            self.errors.append(f"npm: {e}")

    def _api_stackoverflow(self, name_or_username: str):
        self._log(f"   📚 Stack Overflow: {name_or_username}")
        try:
            r = requests.get(
                f"https://api.stackexchange.com/2.3/users?order=desc&sort=reputation"
                f"&inname={quote_plus(name_or_username)}&site=stackoverflow&pagesize=3",
                headers=HEADERS, timeout=TIMEOUT)
            if r.status_code != 200:
                return
            items = r.json().get("items", [])
            if not items:
                return
            u = items[0]
            # Get top answers
            answers = []
            ra = requests.get(
                f"https://api.stackexchange.com/2.3/users/{u['user_id']}/answers"
                f"?order=desc&sort=votes&site=stackoverflow&pagesize=5&filter=withbody",
                headers=HEADERS, timeout=TIMEOUT)
            if ra.status_code == 200:
                for a in ra.json().get("items", [])[:3]:
                    answers.append({
                        "score":    a.get("score", 0),
                        "accepted": a.get("is_accepted", False),
                        "link":     a.get("link", ""),
                    })
            self.results["api_profiles"]["stackoverflow"] = {
                "found":          True,
                "url":            u.get("link"),
                "display_name":   u.get("display_name"),
                "reputation":     u.get("reputation"),
                "location":       u.get("location"),
                "website":        u.get("website_url"),
                "member_since":   datetime.utcfromtimestamp(
                                      u.get("creation_date", 0)).strftime("%Y-%m-%d"),
                "last_seen":      datetime.utcfromtimestamp(
                                      u.get("last_access_date", 0)).strftime("%Y-%m-%d"),
                "gold_badges":    u.get("badge_counts", {}).get("gold", 0),
                "silver_badges":  u.get("badge_counts", {}).get("silver", 0),
                "bronze_badges":  u.get("badge_counts", {}).get("bronze", 0),
                "top_answers":    answers,
            }
            self._log(f"      ✅ {u.get('display_name')} — {u.get('reputation')} rep")
        except Exception as e:
            self.errors.append(f"Stack Overflow: {e}")

    # ── Email checks ──────────────────────────────────────────────────────────

    def _check_breaches(self, email: str):
        self._log(f"   🔓 Breach check: {email}")
        try:
            # HIBP free endpoint (limited)
            r = requests.get(
                f"https://haveibeenpwned.com/api/v3/breachedaccount/{quote_plus(email)}",
                headers={**HEADERS, "hibp-api-key": "public"},
                timeout=TIMEOUT)
            if r.status_code == 200:
                breaches = r.json()
                self.results["data_breaches"] = {
                    "found":    True,
                    "count":    len(breaches),
                    "breaches": [{
                        "name":         b.get("Name"),
                        "domain":       b.get("Domain"),
                        "breach_date":  b.get("BreachDate"),
                        "pwn_count":    b.get("PwnCount", 0),
                        "data_classes": b.get("DataClasses", []),
                        "description":  re.sub(r'<[^>]+>', '', b.get("Description", ""))[:300],
                        "verified":     b.get("IsVerified", False),
                        "sensitive":    b.get("IsSensitive", False),
                    } for b in breaches],
                }
                self._log(f"      ⚠️  In {len(breaches)} breaches")
            elif r.status_code == 404:
                self.results["data_breaches"] = {"found": False}
                self._log(f"      ✅ Not found in breaches")
            else:
                self.results["data_breaches"] = {
                    "found": None,
                    "message": "Check manually: haveibeenpwned.com (API key required for full access)"
                }
        except Exception as e:
            self.errors.append(f"HIBP: {e}")

    def _check_gravatar(self, email: str):
        self._log(f"   🖼️  Gravatar: {email}")
        try:
            h = hashlib.md5(email.strip().lower().encode()).hexdigest()
            r = requests.get(f"https://www.gravatar.com/{h}.json",
                             headers=HEADERS, timeout=TIMEOUT)
            if r.status_code != 200:
                return
            entry = r.json().get("entry", [{}])[0]
            self.results["api_profiles"]["gravatar"] = {
                "found":        True,
                "display_name": entry.get("displayName"),
                "real_name":    entry.get("name", {}).get("formatted"),
                "location":     entry.get("currentLocation"),
                "about":        entry.get("aboutMe", "")[:300],
                "profile_url":  entry.get("profileUrl"),
                "avatar":       f"https://www.gravatar.com/avatar/{h}?s=200",
                "accounts":     [{"domain": a.get("domain"), "url": a.get("url")}
                                  for a in entry.get("accounts", [])],
                "urls":         [{"title": u.get("title"), "value": u.get("value")}
                                  for u in entry.get("urls", [])],
            }
            self._log(f"      ✅ Profile found: {entry.get('displayName')}")
        except Exception as e:
            self.errors.append(f"Gravatar: {e}")

    def _check_email_format(self, email: str):
        """Extract provider info and check for common patterns."""
        domain = email.split("@")[-1].lower() if "@" in email else ""
        provider_map = {
            "gmail.com": "Google", "yahoo.com": "Yahoo", "hotmail.com": "Microsoft",
            "outlook.com": "Microsoft", "protonmail.com": "ProtonMail (privacy-focused)",
            "icloud.com": "Apple", "me.com": "Apple", "aol.com": "AOL",
        }
        self.results["email_info"] = {
            "address":  email,
            "domain":   domain,
            "provider": provider_map.get(domain, f"Custom domain ({domain})"),
            "is_disposable": domain in {
                "mailinator.com", "guerrillamail.com", "tempmail.com",
                "throwaway.email", "yopmail.com", "sharklasers.com",
            },
        }

    # ── Phone lookup ──────────────────────────────────────────────────────────

    def _check_phone(self, phone: str):
        self._log(f"   📞 Phone analysis: {phone}")
        try:
            # Clean number
            digits = re.sub(r'\D', '', phone)
            country = "Unknown"
            carrier = "Unknown"

            # Basic country code detection
            if digits.startswith("1") and len(digits) == 11:
                country = "USA/Canada"
            elif digits.startswith("44"):
                country = "United Kingdom"
            elif digits.startswith("61"):
                country = "Australia"
            elif digits.startswith("91"):
                country = "India"
            elif digits.startswith("49"):
                country = "Germany"
            elif digits.startswith("33"):
                country = "France"

            self.results["phone_info"] = {
                "number":       phone,
                "digits":       digits,
                "country":      country,
                "search_links": [
                    f"https://www.truecaller.com/search/us/{digits}",
                    f"https://www.whitepages.com/phone/{digits}",
                    f"https://www.spokeo.com/phone-search?q={digits}",
                ]
            }
            self._log(f"      ✅ Country: {country}")
        except Exception as e:
            self.errors.append(f"Phone: {e}")

    # ── Domain/website check ──────────────────────────────────────────────────

    def _check_domain(self, website: str):
        self._log(f"   🌐 Domain info: {website}")
        try:
            domain = re.sub(r'^https?://', '', website).split('/')[0]
            # WHOIS via public API
            r = requests.get(
                f"https://api.domainsdb.info/v1/domains/search?domain={domain}&zone=com",
                headers=HEADERS, timeout=TIMEOUT)
            info = {"domain": domain, "url": website}
            if r.status_code == 200:
                data = r.json().get("domains", [])
                if data:
                    d = data[0]
                    info.update({
                        "created":  d.get("create_date", "")[:10],
                        "updated":  d.get("update_date", "")[:10],
                    })

            # Also try to get meta info from the page itself
            try:
                pr = requests.get(website if website.startswith("http") else f"https://{website}",
                                  headers=HEADERS, timeout=8)
                if pr.status_code == 200:
                    title_m = re.search(r'<title[^>]*>(.*?)</title>', pr.text, re.I | re.S)
                    desc_m  = re.search(r'<meta[^>]+name=["\']description["\'][^>]+content=["\'](.*?)["\']',
                                        pr.text, re.I)
                    info["page_title"]   = title_m.group(1).strip() if title_m else None
                    info["page_desc"]    = desc_m.group(1).strip()[:200] if desc_m else None
                    info["server"]       = pr.headers.get("Server")
                    info["technologies"] = []
                    tech_hints = {
                        "WordPress": "wp-content", "Shopify": "shopify",
                        "React": "__NEXT_DATA__", "Angular": "ng-version",
                        "Django": "csrfmiddlewaretoken", "Laravel": "laravel",
                    }
                    for tech, hint in tech_hints.items():
                        if hint in pr.text:
                            info["technologies"].append(tech)
            except Exception:
                pass

            self.results["domain_info"] = info
            self._log(f"      ✅ Domain analyzed")
        except Exception as e:
            self.errors.append(f"Domain: {e}")

    # ── Deep web search ───────────────────────────────────────────────────────

    def _deep_web_search(self, name, username, email, location,
                          employer, phone, website):
        if not self._search:
            self.results["web_presence"]["note"] = "Web search unavailable"
            return

        self._log("   🔍 Running deep web searches...")
        queries = []
        base_name = f'"{name}"' if name else None
        loc_str   = f' "{location}"' if location else ""
        emp_str   = f' "{employer}"' if employer else ""

        if name:
            queries += [
                # Identity
                (f'{base_name}{loc_str}',                        "General mentions"),
                (f'{base_name}{emp_str}',                        "Work mentions"),
                (f'{base_name} email OR contact',                "Contact info"),
                (f'{base_name} phone OR mobile',                 "Phone mentions"),
                # Social
                (f'{base_name} site:linkedin.com',               "LinkedIn"),
                (f'{base_name} site:twitter.com OR x.com',       "Twitter/X"),
                (f'{base_name} site:instagram.com',              "Instagram"),
                (f'{base_name} site:facebook.com',               "Facebook"),
                (f'{base_name} site:tiktok.com',                 "TikTok"),
                (f'{base_name} site:youtube.com',                "YouTube"),
                # Professional
                (f'{base_name} resume OR CV OR portfolio',       "Resume/Portfolio"),
                (f'{base_name} site:medium.com OR substack.com', "Articles/Blog"),
                (f'{base_name} site:github.com',                 "GitHub"),
                # News
                (f'{base_name} site:news.google.com OR site:reuters.com', "News mentions"),
                (f'{base_name} arrested OR charged OR convicted', "Legal records"),
                (f'{base_name} married OR wedding OR spouse',    "Relationship info"),
                (f'{base_name} address OR lives in OR located',  "Location info"),
            ]

        if username:
            queries += [
                (f'"{username}" profile',                        "Username profile"),
                (f'"{username}" site:reddit.com',                "Reddit"),
                (f'"{username}" site:twitch.tv',                 "Twitch"),
                (f'"{username}" site:steam community.com',       "Steam"),
                (f'"{username}" site:pastebin.com',              "Pastebin"),
                (f'"{username}" discord',                        "Discord mentions"),
            ]

        if email:
            queries += [
                (f'"{email}"',                                   "Email mentions"),
                (f'"{email}" site:pastebin.com',                 "Email in pastes"),
            ]

        if phone:
            queries += [
                (f'"{phone}"',                                   "Phone mentions"),
            ]

        results = {}
        for query, label in queries:
            try:
                result = self._search(query)
                if result and "No results" not in result and len(result) > 60:
                    results[label] = result[:800]
                    self._log(f"      ✅ {label}")
                time.sleep(0.3)
            except Exception as e:
                self.errors.append(f"Search [{label}]: {e}")

        self.results["web_presence"]["search_results"] = results
        self._log(f"   ✅ {len(results)}/{len(queries)} searches returned results")

    # ── Image search links ────────────────────────────────────────────────────

    def _generate_image_search_links(self, name: str, location: str = None):
        """Generate links to reverse image / person search engines."""
        q = quote_plus(name + (f" {location}" if location else ""))
        self.results["image_search"] = {
            "google_images":   f"https://www.google.com/search?q={q}&tbm=isch",
            "google_people":   f"https://www.google.com/search?q={q}",
            "bing_images":     f"https://www.bing.com/images/search?q={q}",
            "pimeyes":         f"https://pimeyes.com/en",
            "tineye":          f"https://tineye.com",
            "spokeo":          f"https://www.spokeo.com/search?q={q}",
            "intelius":        f"https://www.intelius.com/search/?firstname={quote_plus(name.split()[0])}&lastname={quote_plus(name.split()[-1])}",
            "beenverified":    f"https://www.beenverified.com/people/{q}/",
            "fastpeoplesearch":f"https://www.fastpeoplesearch.com/name/{q}",
        }

    # ── Summary ───────────────────────────────────────────────────────────────

    def _build_summary(self):
        r = self.results
        summary = []

        # Platform hits
        hits = r.get("platform_hits", [])
        if hits:
            summary.append(f"✅ USERNAME FOUND ON {len(hits)} PLATFORMS")
            for h in hits[:10]:
                summary.append(f"   • {h['platform']}: {h['url']}")

        # API profiles
        for platform, data in r.get("api_profiles", {}).items():
            if data.get("found"):
                summary.append(f"✅ {platform.upper()} PROFILE: {data.get('url', '')}")

        # Breaches
        b = r.get("data_breaches", {})
        if b.get("found"):
            summary.append(f"⚠️  DATA BREACHES: Found in {b['count']} breach(es)")
        elif b.get("found") is False:
            summary.append("✅ NO DATA BREACHES found for this email")

        # Web
        wp = len(r.get("web_presence", {}).get("search_results", {}))
        if wp:
            summary.append(f"🌐 WEB PRESENCE: {wp} search categories returned results")

        # Phone / domain
        if r.get("phone_info", {}).get("country"):
            summary.append(f"📞 PHONE: {r['phone_info']['country']}")
        if r.get("domain_info", {}).get("domain"):
            summary.append(f"🌐 DOMAIN: {r['domain_info']['domain']} analyzed")

        if self.errors:
            summary.append(f"⚠️  {len(self.errors)} sources had errors")

        r["summary"] = summary
        r["errors"]  = self.errors


# ═══════════════════════════════════════════════════════════
# REPORT GENERATOR
# ═══════════════════════════════════════════════════════════

def generate_report(results: dict, output_path: str = None) -> str:
    query     = results.get("query", {})
    subject   = query.get("name") or query.get("username") or query.get("email") or "Unknown"
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")

    if not output_path:
        safe = re.sub(r'[^a-zA-Z0-9]', '_', subject)
        output_path = f"osint_report_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        def h(text, level=1):
            p = doc.add_heading(text, level=level)
            return p

        def p(text, bold=False, italic=False):
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold, run.italic = bold, italic
            return para

        def kv(key, val):
            if not val:
                return
            para = doc.add_paragraph()
            run_k = para.add_run(f"{key}: ")
            run_k.bold = True
            para.add_run(str(val))

        # Title
        title = doc.add_heading(f"OSINT Intelligence Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(f"Subject: {subject}  |  Generated: {timestamp}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].italic = True
        doc.add_paragraph()

        # Summary
        h("Executive Summary")
        for line in results.get("summary", []):
            doc.add_paragraph(line)
        doc.add_paragraph()

        # Query params
        h("Search Parameters")
        for k, v in query.items():
            if v and k != "timestamp":
                kv(k.title(), v)
        doc.add_paragraph()

        # Platform hits
        hits = results.get("platform_hits", [])
        if hits:
            h("Platform Presence")
            p(f"Username found on {len(hits)} platforms:", bold=True)
            for hit in hits:
                doc.add_paragraph(f"• {hit['platform']}: {hit['url']}", style="List Bullet")
            doc.add_paragraph()

        # API profiles
        for platform, data in results.get("api_profiles", {}).items():
            if not data.get("found"):
                continue
            h(platform.upper(), level=2)
            skip = {"found", "top_repos", "verified_proofs", "accounts", "urls",
                    "top_articles", "recent_posts", "top_subreddits", "recent_activity",
                    "top_answers", "packages", "top_languages", "organizations",
                    "starred_repos", "projects", "search_links"}
            for k, v in data.items():
                if k not in skip and v is not None and v != "" and v != [] and v != {}:
                    kv(k.replace("_", " ").title(), v)

            # Render lists
            for list_key, label in [
                ("top_repos",       "Top Repositories"),
                ("top_languages",   "Top Languages"),
                ("organizations",   "Organizations"),
                ("verified_proofs", "Verified Identities"),
                ("top_subreddits",  "Active Subreddits"),
                ("top_articles",    "Articles"),
                ("recent_posts",    "Recent Posts"),
                ("packages",        "npm Packages"),
                ("projects",        "Projects"),
            ]:
                items = data.get(list_key)
                if items:
                    p(f"\n{label}:", bold=True)
                    for item in (items[:8] if isinstance(items, list) else []):
                        if isinstance(item, dict):
                            line = " | ".join(f"{k}: {v}" for k, v in item.items()
                                              if v and k not in ("url", "links"))
                            doc.add_paragraph(f"• {line}", style="List Bullet")
                        else:
                            doc.add_paragraph(f"• {item}", style="List Bullet")
            doc.add_paragraph()

        # Email info
        ei = results.get("email_info")
        if ei:
            h("Email Analysis", level=2)
            for k, v in ei.items():
                kv(k.replace("_", " ").title(), v)
            doc.add_paragraph()

        # Breaches
        b = results.get("data_breaches", {})
        h("Data Breach Analysis")
        if b.get("found"):
            p(f"⚠ Found in {b['count']} known breach(es)", bold=True)
            for breach in b.get("breaches", []):
                h(breach.get("name", "Unknown"), level=2)
                kv("Domain",       breach.get("domain"))
                kv("Date",         breach.get("breach_date"))
                kv("Records",      f"{breach.get('pwn_count', 0):,}")
                kv("Data Exposed", ", ".join(breach.get("data_classes", [])))
                kv("Verified",     breach.get("verified"))
                if breach.get("description"):
                    p(breach["description"])
        else:
            p(b.get("message", "No breach data for this email."))
        doc.add_paragraph()

        # Phone
        pi = results.get("phone_info")
        if pi:
            h("Phone Analysis", level=2)
            kv("Number",  pi.get("number"))
            kv("Country", pi.get("country"))
            if pi.get("search_links"):
                p("Public lookup links:", bold=True)
                for link in pi["search_links"]:
                    doc.add_paragraph(f"• {link}", style="List Bullet")
            doc.add_paragraph()

        # Domain
        di = results.get("domain_info")
        if di:
            h("Website / Domain Analysis", level=2)
            for k, v in di.items():
                if k != "technologies" and v:
                    kv(k.replace("_", " ").title(), v)
            if di.get("technologies"):
                kv("Technologies", ", ".join(di["technologies"]))
            doc.add_paragraph()

        # Image/person search links
        img = results.get("image_search", {})
        if img:
            h("Person & Image Search Links")
            p("Use these to manually verify photos and find more information:", italic=True)
            for label, url in img.items():
                doc.add_paragraph(f"• {label.replace('_', ' ').title()}: {url}",
                                   style="List Bullet")
            doc.add_paragraph()

        # Web presence
        web = results.get("web_presence", {}).get("search_results", {})
        if web:
            h("Web Presence")
            for category, content in web.items():
                h(category, level=2)
                doc.add_paragraph(content[:1000])
            doc.add_paragraph()

        # Errors
        if results.get("errors"):
            h("Scan Errors / Skipped Sources")
            for err in results["errors"]:
                doc.add_paragraph(f"• {err}", style="List Bullet")

        doc.save(output_path)
        return output_path

    except ImportError:
        txt_path = output_path.replace(".docx", ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"OSINT REPORT: {subject}\nGenerated: {timestamp}\n")
            f.write("=" * 60 + "\n")
            for line in results.get("summary", []):
                f.write(line + "\n")
        return txt_path