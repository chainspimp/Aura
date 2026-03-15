# ============================================
# FILE: ai/agent.py
# Autonomous Task Chain Executor
# Wires planner.py to real tools and executes
# multi-step plans without user prompting.
# ============================================

import os
import re
import time
import logging
import requests
from datetime import datetime
from typing import List, Dict

from config import OLLAMA_API_URL, OLLAMA_MODEL
from ai.planner import create_plan

logger = logging.getLogger(__name__)

TOOL_WEB_SEARCH    = "web_search"
TOOL_DEEP_RESEARCH = "deep_research"
TOOL_THINKING      = "thinking"
TOOL_CODE          = "code_generation"
TOOL_IMAGE         = "image_generation"
TOOL_VISION        = "vision"
TOOL_SAVE_DOC      = "save_doc"   # saves a .docx Word report
TOOL_SAVE_TXT      = "save_txt"   # saves a plain .txt file
TOOL_FINAL         = "final_answer"


class AgentStep:
    def __init__(self, tool: str, input_text: str):
        self.tool = tool
        self.input = input_text
        self.output = ""
        self.success = False
        self.duration = 0.0


class AgentResult:
    def __init__(self):
        self.steps: List[AgentStep] = []
        self.final_answer: str = ""
        self.saved_files: List[str] = []
        self.success: bool = False
        self.total_time: float = 0.0

    def summary(self) -> str:
        lines = [f"🤖 Agent completed {len(self.steps)} steps in {self.total_time:.1f}s"]
        for i, s in enumerate(self.steps, 1):
            status = "✅" if s.success else "❌"
            lines.append(f"  {i}. {status} {s.tool} ({s.duration:.1f}s)")
        if self.saved_files:
            lines.append("\n📁 Files saved:")
            for f in self.saved_files:
                lines.append(f"   • {f}")
        return "\n".join(lines)


class AutonomousAgent:
    """
    Executes multi-step task chains planned by the LLM.

    Usage:
        agent = AutonomousAgent(tool_executor, thinking_system, coding_system)
        result = agent.run("Research AI trends and write a Word report")
        print(result.final_answer)
    """

    def __init__(self, tool_executor, thinking_system=None, coding_system=None):
        self.tools = tool_executor
        self.thinking = thinking_system
        self.coding = coding_system
        self.output_dir = "agent_outputs"
        os.makedirs(self.output_dir, exist_ok=True)

    # ── Public entry point ────────────────────────────────────────────────────

    def run(self, prompt: str, context: str = "", on_step=None) -> AgentResult:
        """
        Run a full autonomous task chain.

        Args:
            prompt:   The user's high-level request
            context:  Optional memory context to include
            on_step:  Optional callback(step_num, tool, msg) for live UI updates

        Returns:
            AgentResult with steps, final answer, and saved file paths
        """
        result = AgentResult()
        start_total = time.time()

        print(f"\n{'='*60}")
        print(f"🤖 AGENT MODE — Planning task...")
        print(f"{'='*60}")

        # 1. Build the plan
        try:
            steps = create_plan(prompt, context)
        except Exception as e:
            logger.error(f"Planning failed: {e}")
            result.final_answer = f"I couldn't create a plan for that: {e}"
            return result

        print(f"📋 Plan ({len(steps)} steps):")
        for i, s in enumerate(steps, 1):
            print(f"   {i}. [{s['tool']}] {s['input'][:70]}")
        print()

        # 2. Execute each step, accumulating context
        accumulated_context = context
        step_outputs: Dict[str, str] = {}

        for i, step_def in enumerate(steps, 1):
            tool = step_def.get("tool", TOOL_FINAL)
            inp  = step_def.get("input", prompt)
            inp  = self._inject_context(inp, step_outputs)

            step = AgentStep(tool, inp)
            step_start = time.time()

            if on_step:
                on_step(i, tool, f"Step {i}/{len(steps)}: {tool}")

            print(f"⚡ Step {i}/{len(steps)}: [{tool}] — {inp[:60]}")

            try:
                output = self._execute_step(tool, inp, accumulated_context, result)
                step.output = output
                step.success = True
                accumulated_context += f"\n\n[Step {i} — {tool}]\n{output}"
                step_outputs[tool] = output
            except Exception as e:
                logger.error(f"Step {i} ({tool}) failed: {e}")
                step.output = f"Error: {e}"
                step.success = False

            step.duration = time.time() - step_start
            result.steps.append(step)
            print(f"   ✓ Done ({step.duration:.1f}s)\n")

            if tool == TOOL_FINAL:
                result.final_answer = step.output
                break

        # 3. Synthesise final answer if no explicit final_answer step ran
        if not result.final_answer:
            result.final_answer = self._synthesise_answer(prompt, accumulated_context)

        result.success = all(s.success for s in result.steps)
        result.total_time = time.time() - start_total
        return result

    # ── Step dispatcher ───────────────────────────────────────────────────────

    def _execute_step(self, tool: str, inp: str, context: str, result: AgentResult) -> str:
        if tool == TOOL_WEB_SEARCH:
            return self.tools.web_search(inp)

        elif tool == TOOL_DEEP_RESEARCH:
            return self.tools.deep_research(inp)

        elif tool == TOOL_THINKING:
            if self.thinking:
                _, conclusion = self.thinking.deep_think(inp, context)
                return conclusion
            return self._llm_step(inp, context)

        elif tool == TOOL_CODE:
            if self.coding:
                res = self.coding.generate_and_save(inp, context)
                if res.get('success'):
                    result.saved_files.append(res['filepath'])
                    return (f"Code generated: {res['filename']} "
                            f"({res['language']}, {res['total_lines']} lines)\n"
                            f"Saved to: {res['filepath']}")
                return f"Code generation failed: {res.get('error')}"
            return "Code generation unavailable."

        elif tool == TOOL_IMAGE:
            res, path = self.tools.generate_image_local(inp)
            if path:
                result.saved_files.append(path)
            return res

        elif tool == TOOL_VISION:
            from ai.vision import get_visual_context
            return get_visual_context(force=True)

        elif tool == TOOL_SAVE_DOC:
            return self._save_word_doc(inp, context, result)

        elif tool == TOOL_SAVE_TXT:
            return self._save_txt(inp, context, result)

        elif tool == TOOL_FINAL:
            return self._synthesise_answer(inp, context)

        else:
            logger.warning(f"Unknown tool '{tool}', falling back to LLM")
            return self._llm_step(inp, context)

    # ── Tool implementations ──────────────────────────────────────────────────

    def _llm_step(self, task: str, context: str) -> str:
        """Generic LLM reasoning step."""
        prompt = (
            "You are AURA, an expert AI assistant.\n\n"
            f"Context from previous steps:\n{context}\n\n"
            f"Current task: {task}\n\n"
            "Complete this task thoroughly and concisely."
        )
        try:
            resp = requests.post(
                OLLAMA_API_URL,
                json={
                    "model": OLLAMA_MODEL,
                    "prompt": prompt,
                    "stream": False,
                    "options": {"temperature": 0.7, "num_predict": 800}
                },
                timeout=120
            )
            resp.raise_for_status()
            return resp.json().get("response", "").strip()
        except Exception as e:
            return f"LLM step failed: {e}"

    def _synthesise_answer(self, original_prompt: str, accumulated_context: str) -> str:
        """Synthesise a coherent final answer from all accumulated step outputs."""
        prompt = (
            "You are AURA. You have just completed a multi-step autonomous task.\n\n"
            f"Original request: {original_prompt}\n\n"
            f"All work completed:\n{accumulated_context}\n\n"
            "Write a clear, complete, well-structured final answer that directly addresses "
            "the original request using the information gathered above."
        )
        try:
            resp = requests.post(
                OLLAMA_API_URL,
                json={
                    "model": OLLAMA_MODEL,
                    "prompt": prompt,
                    "stream": False,
                    "options": {"temperature": 0.6, "num_predict": 1000}
                },
                timeout=180
            )
            resp.raise_for_status()
            return resp.json().get("response", "").strip()
        except Exception as e:
            return f"Could not synthesise final answer: {e}"

    def _save_word_doc(self, content_prompt: str, context: str, result: AgentResult) -> str:
        """Generate and save a Word .docx report. Falls back to .txt if python-docx missing."""
        # Generate the actual report content
        doc_content = self._llm_step(
            f"Write a complete, well-structured report for: {content_prompt}\n"
            "Use ## for section headings and write in clear paragraphs.",
            context
        )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name  = re.sub(r'[^a-zA-Z0-9]', '_', content_prompt[:35]).strip('_')

        try:
            from docx import Document as DocxDocument
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = DocxDocument()

            # Title + date
            title = doc.add_heading(safe_name.replace('_', ' ').title(), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_p = doc.add_paragraph(
                f"Generated by AURA — {datetime.now().strftime('%B %d, %Y')}"
            )
            date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # Parse markdown-ish headings and body text
            for line in doc_content.split('\n'):
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith(('- ', '* ')):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)

            filepath = os.path.join(self.output_dir, f"{safe_name}_{timestamp}.docx")
            doc.save(filepath)
            result.saved_files.append(filepath)
            return f"Word document saved: {filepath}"

        except ImportError:
            logger.warning("python-docx not installed — saving as .txt instead. "
                           "Run: pip install python-docx")
            return self._save_txt(doc_content, context, result,
                                  suffix=timestamp, name=safe_name)
        except Exception as e:
            logger.error(f"Word doc save error: {e}")
            return self._save_txt(doc_content, context, result,
                                  suffix=timestamp, name=safe_name)

    def _save_txt(self, content: str, context: str, result: AgentResult,
                  suffix: str = None, name: str = None) -> str:
        """Save output as a plain text file."""
        if not suffix:
            suffix = datetime.now().strftime("%Y%m%d_%H%M%S")
        if not name:
            name = "agent_output"

        # If content is too short it's probably a task description — generate content first
        if len(content) < 100 and '\n' not in content:
            content = self._llm_step(content, context)

        filepath = os.path.join(self.output_dir, f"{name}_{suffix}.txt")
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            result.saved_files.append(filepath)
            return f"Text file saved: {filepath}"
        except Exception as e:
            return f"File save failed: {e}"

    # ── Helper ────────────────────────────────────────────────────────────────

    def _inject_context(self, inp: str, step_outputs: Dict[str, str]) -> str:
        """Replace {{tool_name}} placeholders with previous step outputs."""
        for tool, output in step_outputs.items():
            placeholder = f"{{{{{tool}}}}}"
            if placeholder in inp:
                inp = inp.replace(placeholder, output[:500])
        return inp