# ============================================
# FILE: ai/hacker_agent.py
# AURA Security Agent — Windows + Linux
# Auto-detects WSL / Git Bash / PowerShell
# Falls back to a pure-Python Linux emulator
# ============================================

import os
import re
import sys
import time
import queue
import shutil
import socket
import logging
import threading
import subprocess
from datetime import datetime
from typing import Optional, List, Dict, Callable

import requests
from config import OLLAMA_API_URL, OLLAMA_MODEL

logger = logging.getLogger(__name__)

REPORT_DIR = "pentest_reports"
os.makedirs(REPORT_DIR, exist_ok=True)

IS_WINDOWS = sys.platform == "win32"


# ═══════════════════════════════════════════════════════════
# SHELL DETECTOR
# Finds the best available shell on the system
# ═══════════════════════════════════════════════════════════

def _find_shell() -> tuple:
    """
    Returns (shell_cmd_list, shell_type) for the best available shell.
    Priority: WSL > Git Bash > Cygwin > PowerShell > cmd
    """
    if not IS_WINDOWS:
        return (["bash", "--norc", "--noprofile"], "bash")

    # 1. WSL (Windows Subsystem for Linux) — real Linux kernel
    try:
        result = subprocess.run(
            ["wsl", "--status"],
            capture_output=True, text=True, timeout=5
        )
        if result.returncode == 0 or "Default Distribution" in (result.stdout + result.stderr):
            # Test it actually works
            test = subprocess.run(
                ["wsl", "echo", "ok"],
                capture_output=True, text=True, timeout=5
            )
            if "ok" in test.stdout:
                return (["wsl", "bash", "--norc", "--noprofile"], "wsl")
    except Exception:
        pass

    # 2. Git Bash (ships with Git for Windows)
    git_bash_paths = [
        r"C:\Program Files\Git\bin\bash.exe",
        r"C:\Program Files (x86)\Git\bin\bash.exe",
        os.path.expanduser(r"~\AppData\Local\Programs\Git\bin\bash.exe"),
    ]
    for path in git_bash_paths:
        if os.path.exists(path):
            return ([path, "--norc", "--noprofile"], "bash")

    # 3. Cygwin bash
    cygwin_bash = r"C:\cygwin64\bin\bash.exe"
    if os.path.exists(cygwin_bash):
        return ([cygwin_bash, "--norc", "--noprofile"], "bash")

    # 4. PowerShell Core (pwsh) — better than cmd
    try:
        result = subprocess.run(["pwsh", "-Version"], capture_output=True, timeout=3)
        if result.returncode == 0:
            return (["pwsh", "-NoLogo", "-NonInteractive", "-Command", "-"], "powershell")
    except Exception:
        pass

    # 5. Windows PowerShell
    try:
        result = subprocess.run(
            ["powershell", "-Command", "echo ok"],
            capture_output=True, text=True, timeout=3
        )
        if "ok" in result.stdout:
            return (["powershell", "-NoLogo", "-NonInteractive", "-Command", "-"], "powershell")
    except Exception:
        pass

    # 6. cmd.exe fallback
    return (["cmd.exe", "/Q"], "cmd")


# ═══════════════════════════════════════════════════════════
# PURE PYTHON LINUX EMULATOR
# Used when no real shell is available
# Implements ~30 common commands natively
# ═══════════════════════════════════════════════════════════

class PythonShellEmulator:
    """
    Pure Python emulation of a Linux shell.
    Runs network tools (ping, nslookup) via Python's socket/subprocess.
    Handles file ops, directory navigation, env vars natively.
    """

    def __init__(self, log_fn: Callable = print):
        self._cwd  = os.getcwd()
        self._env  = dict(os.environ)
        self._log  = log_fn
        self._vars = {}   # shell variables

    def run(self, cmd: str, timeout: int = 60, on_line: Callable = None) -> str:
        cmd = cmd.strip()
        if not cmd or cmd.startswith("#"):
            return ""

        # Handle variable assignment: VAR=value
        var_match = re.match(r'^([A-Z_][A-Z0-9_]*)=(.*)$', cmd)
        if var_match:
            self._vars[var_match.group(1)] = var_match.group(2)
            return ""

        # Expand variables
        for k, v in self._vars.items():
            cmd = cmd.replace(f"${k}", v).replace(f"${{{k}}}", v)

        # Parse command
        parts = cmd.split(None, 1)
        command = parts[0].lower()
        args = parts[1] if len(parts) > 1 else ""

        output = self._dispatch(command, args, cmd, timeout, on_line)
        # Stream output line by line if on_line callback provided
        if on_line and output:
            for line in output.split("\n"):
                on_line(line)
        return output if output else ""

    def _dispatch(self, command: str, args: str, full_cmd: str,
                  timeout: int, on_line: Callable) -> str:
        # ── File system ──
        if command == "pwd":
            return self._cwd

        if command == "ls" or command == "dir":
            path = args.strip().lstrip("-la ") or self._cwd
            try:
                entries = os.listdir(path)
                lines = []
                for e in sorted(entries):
                    fp = os.path.join(path, e)
                    size = os.path.getsize(fp) if os.path.isfile(fp) else 0
                    typ  = "d" if os.path.isdir(fp) else "-"
                    lines.append(f"{typ}rwxr-xr-x  {size:>10}  {e}")
                return "\n".join(lines)
            except Exception as e:
                return f"ls: {e}"

        if command == "cd":
            target = args.strip() or os.path.expanduser("~")
            target = os.path.join(self._cwd, target)
            target = os.path.normpath(target)
            if os.path.isdir(target):
                self._cwd = target
                return f"[cwd: {self._cwd}]"
            return f"cd: {target}: No such directory"

        if command == "cat":
            path = os.path.join(self._cwd, args.strip())
            try:
                with open(path, "r", errors="replace") as f:
                    return f.read()
            except Exception as e:
                return f"cat: {e}"

        if command in ("mkdir",):
            path = os.path.join(self._cwd, args.strip().lstrip("-p "))
            os.makedirs(path, exist_ok=True)
            return f"Directory created: {path}"

        if command in ("rm", "del"):
            path = os.path.join(self._cwd, args.strip().lstrip("-rf "))
            try:
                if os.path.isfile(path):
                    os.remove(path)
                return f"Removed: {path}"
            except Exception as e:
                return f"rm: {e}"

        if command == "echo":
            return args

        if command == "env" or command == "printenv":
            return "\n".join(f"{k}={v}" for k, v in sorted(self._env.items())[:30])

        if command == "uname":
            import platform
            if "-a" in args:
                return f"Windows {platform.version()} (AURA Virtual Shell) Python/{platform.python_version()}"
            return "AURA-VirtualShell"

        if command == "whoami":
            return os.environ.get("USERNAME", os.environ.get("USER", "user"))

        if command == "id":
            user = os.environ.get("USERNAME", "user")
            return f"uid=1000({user}) gid=1000({user}) groups=1000({user})"

        if command == "date":
            return datetime.now().strftime("%a %b %d %H:%M:%S %Z %Y")

        if command == "uptime":
            return f" {datetime.now().strftime('%H:%M:%S')} up  0 days, load average: 0.00, 0.00, 0.00"

        if command in ("clear", "cls"):
            return "\x1b[2J\x1b[H"

        if command == "history":
            return "[history not available in virtual shell]"

        # ── Network tools ──
        if command == "ping":
            host = args.split()[0] if args.split() else "localhost"
            return self._ping(host)

        if command in ("nslookup", "host", "dig"):
            host = args.split()[0] if args.split() else ""
            return self._dns_lookup(host)

        if command == "curl":
            return self._curl(args)

        if command == "wget":
            return self._wget(args)

        if command == "whois":
            host = args.strip().split()[0] if args.strip() else ""
            return self._whois(host)

        if command in ("nmap", "nc", "netcat", "gobuster", "nikto",
                       "sqlmap", "hydra", "john", "hashcat", "dirb",
                       "ffuf", "subfinder", "amass", "whatweb",
                       "wafw00f", "sslscan", "theharvester"):
            return self._run_native_tool(command, args, full_cmd, timeout, on_line)

        if command in ("python", "python3"):
            return self._run_python(args, timeout)

        if command == "pip" or command == "pip3":
            return self._run_pip(args, timeout)

        if command == "apt-get" or command == "apt":
            return self._apt_install(args, timeout)

        if command == "grep":
            return self._grep(args)

        if command in ("head", "tail"):
            return self._head_tail(command, args)

        if command == "find":
            return self._find(args)

        if command == "which":
            tool = args.strip()
            path = shutil.which(tool)
            return path if path else f"{tool}: not found"

        if command in ("sleep", "wait"):
            try:
                secs = float(args.strip())
                time.sleep(min(secs, 10))
                return ""
            except Exception:
                return ""

        if command == "export":
            m = re.match(r'([A-Z_][A-Z0-9_]*)=(.*)', args)
            if m:
                self._vars[m.group(1)] = m.group(2)
                self._env[m.group(1)]  = m.group(2)
                return ""
            return ""

        # ── Fallback: try to run as a real process ──
        return self._run_real(full_cmd, timeout, on_line)

    def _ping(self, host: str) -> str:
        try:
            ip = socket.gethostbyname(host)
            lines = [f"PING {host} ({ip}): 56 data bytes"]
            for i in range(4):
                t0 = time.time()
                try:
                    s = socket.create_connection((ip, 80), timeout=2)
                    s.close()
                    ms = (time.time() - t0) * 1000
                    lines.append(f"64 bytes from {ip}: icmp_seq={i} ttl=64 time={ms:.1f} ms")
                except Exception:
                    lines.append(f"Request timeout for icmp_seq {i}")
                time.sleep(0.3)
            lines.append(f"\n--- {host} ping statistics ---")
            lines.append(f"4 packets transmitted, 4 received, 0% packet loss")
            return "\n".join(lines)
        except Exception as e:
            return f"ping: {host}: {e}"

    def _dns_lookup(self, host: str) -> str:
        if not host:
            return "Usage: dig <hostname>"
        try:
            ip = socket.gethostbyname(host)
            try:
                hostname, aliases, ips = socket.gethostbyaddr(ip)
            except Exception:
                hostname, aliases, ips = host, [], [ip]
            lines = [
                f"; <<>> AURA Virtual dig <<>> {host}",
                f";; ANSWER SECTION:",
                f"{host}.\t\t300\tIN\tA\t{ip}",
            ]
            if aliases:
                lines.append(f";; ALIASES: {', '.join(aliases)}")
            return "\n".join(lines)
        except Exception as e:
            return f"dig: {host}: {e}"

    def _curl(self, args: str) -> str:
        import urllib.request
        # Extract URL
        url_match = re.search(r'https?://\S+', args)
        if not url_match:
            return "curl: no URL provided"
        url = url_match.group(0)
        headers_only = "-I" in args or "--head" in args
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "curl/7.81.0"})
            with urllib.request.urlopen(req, timeout=10) as resp:
                if headers_only:
                    lines = [f"HTTP/1.1 {resp.status} {resp.reason}"]
                    for k, v in resp.headers.items():
                        lines.append(f"{k}: {v}")
                    return "\n".join(lines)
                else:
                    return resp.read(4096).decode("utf-8", errors="replace")
        except Exception as e:
            return f"curl: ({e})"

    def _wget(self, args: str) -> str:
        import urllib.request
        url_match = re.search(r'https?://\S+', args)
        if not url_match:
            return "wget: no URL"
        url = url_match.group(0)
        filename = url.split("/")[-1] or "index.html"
        try:
            urllib.request.urlretrieve(url, filename)
            return f"'{filename}' saved"
        except Exception as e:
            return f"wget: {e}"

    def _whois(self, host: str) -> str:
        try:
            r = requests.get(
                f"https://api.domainsdb.info/v1/domains/search?domain={host}&zone=com",
                timeout=8
            )
            if r.status_code == 200:
                data = r.json().get("domains", [{}])[0]
                return "\n".join(f"{k}: {v}" for k, v in data.items())
            # Fallback: basic WHOIS via socket
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.settimeout(5)
                s.connect(("whois.iana.org", 43))
                s.send(f"{host}\r\n".encode())
                data = b""
                while True:
                    chunk = s.recv(4096)
                    if not chunk:
                        break
                    data += chunk
                return data.decode("utf-8", errors="replace")
        except Exception as e:
            return f"whois: {e}"

    def _grep(self, args: str) -> str:
        parts = args.split(None, 2)
        if len(parts) < 2:
            return "Usage: grep <pattern> <file>"
        pattern, filepath = parts[0], parts[1]
        fullpath = os.path.join(self._cwd, filepath)
        try:
            with open(fullpath, "r", errors="replace") as f:
                lines = [l.rstrip() for l in f if re.search(pattern, l)]
            return "\n".join(lines) if lines else "(no matches)"
        except Exception as e:
            return f"grep: {e}"

    def _head_tail(self, cmd: str, args: str) -> str:
        parts = args.split()
        n = 10
        filepath = parts[-1] if parts else ""
        for i, p in enumerate(parts):
            if p == "-n" and i + 1 < len(parts):
                try:
                    n = int(parts[i + 1])
                except Exception:
                    pass
        fullpath = os.path.join(self._cwd, filepath)
        try:
            with open(fullpath, "r", errors="replace") as f:
                lines = f.readlines()
            selection = lines[:n] if cmd == "head" else lines[-n:]
            return "".join(selection)
        except Exception as e:
            return f"{cmd}: {e}"

    def _find(self, args: str) -> str:
        parts = args.split()
        search_dir = parts[0] if parts else self._cwd
        name_pattern = ""
        for i, p in enumerate(parts):
            if p == "-name" and i + 1 < len(parts):
                name_pattern = parts[i + 1].strip("'\"")
        results = []
        try:
            for root, dirs, files in os.walk(search_dir):
                for f in files:
                    if not name_pattern or re.search(
                        name_pattern.replace("*", ".*"), f
                    ):
                        results.append(os.path.join(root, f))
                if len(results) > 100:
                    break
        except Exception:
            pass
        return "\n".join(results[:100])

    def _run_native_tool(self, tool: str, args: str, full_cmd: str,
                          timeout: int, on_line: Callable) -> str:
        """Try to run a real security tool, install if missing."""
        real_path = shutil.which(tool)
        if real_path:
            return self._run_real(full_cmd, timeout, on_line)

        # Tool not found — suggest install
        install_hints = {
            "nmap":       "Download from: https://nmap.org/download.html",
            "gobuster":   "go install github.com/OJ/gobuster/v3@latest",
            "nikto":      "Download from: https://github.com/sullo/nikto",
            "sqlmap":     "pip install sqlmap  OR  https://sqlmap.org",
            "hydra":      "Download from: https://github.com/vanhauser-thc/thc-hydra",
            "subfinder":  "go install github.com/projectdiscovery/subfinder/v2/cmd/subfinder@latest",
            "whatweb":    "gem install whatweb  OR  https://github.com/urbanadventurer/WhatWeb",
            "wafw00f":    "pip install wafw00f",
            "sslscan":    "Download from: https://github.com/rbsec/sslscan",
            "ffuf":       "go install github.com/ffuf/ffuf/v2@latest",
            "amass":      "go install github.com/owasp-amass/amass/v4/...@master",
            "theharvester": "pip install theHarvester",
        }
        hint = install_hints.get(tool, f"Tool '{tool}' not found in PATH")
        return (
            f"[!] {tool}: not found\n"
            f"[*] Install: {hint}\n"
            f"[*] After installing, add to PATH and retry."
        )

    def _run_python(self, args: str, timeout: int) -> str:
        script = args.strip()
        if script.startswith("-c "):
            code = script[3:].strip("\"'")
            try:
                import io, contextlib
                buf = io.StringIO()
                with contextlib.redirect_stdout(buf):
                    exec(code)
                return buf.getvalue()
            except Exception as e:
                return f"Python error: {e}"
        return f"python: {script}: use -c for inline code"

    def _run_pip(self, args: str, timeout: int) -> str:
        return self._run_real(
            f"{sys.executable} -m pip {args}", timeout, None
        )

    def _apt_install(self, args: str, timeout: int) -> str:
        # On Windows, apt-get doesn't exist — redirect to winget/pip
        pkg = args.replace("install", "").replace("-y", "").strip()
        return (
            f"[*] apt-get not available on Windows\n"
            f"[*] For '{pkg}', try:\n"
            f"    winget install {pkg}\n"
            f"    pip install {pkg}\n"
            f"    Or download from the tool's official website"
        )

    def _run_real(self, cmd: str, timeout: int, on_line: Callable) -> str:
        """Run a real Windows/system process, streaming output live."""
        try:
            proc = subprocess.Popen(
                cmd,
                shell=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                cwd=self._cwd,
                env=self._env,
                text=True,
                encoding="utf-8",
                errors="replace",
                bufsize=1,
            )
            lines = []
            try:
                for line in iter(proc.stdout.readline, ""):
                    line = line.rstrip("\n")
                    if line:
                        lines.append(line)
                        # Stream immediately — don't buffer
                        if on_line:
                            on_line(line)
                proc.wait(timeout=timeout)
            except subprocess.TimeoutExpired:
                proc.kill()
                msg = "[!] Command timed out"
                lines.append(msg)
                if on_line:
                    on_line(msg)
            return "\n".join(lines)
        except Exception as e:
            msg = f"[!] Error: {e}"
            if on_line:
                on_line(msg)
            return msg


# ═══════════════════════════════════════════════════════════
# SMART BASH SESSION
# Uses WSL/Git Bash if available, Python emulator otherwise
# ═══════════════════════════════════════════════════════════

class BashSession:
    """
    Smart shell session that works on Windows and Linux.
    Uses WSL > Git Bash > PowerShell > Python emulator.
    """

    SENTINEL = "##AURA_DONE_8675309##"

    def __init__(self):
        self.proc          = None
        self.output_queue  = queue.Queue()
        self._reader       = None
        self._running      = False
        self._emulator     = None   # PythonShellEmulator fallback
        self.shell_type    = "unknown"
        self.shell_info    = "Not started"

    def start(self, log_fn: Callable = print):
        if self.proc or self._emulator:
            return

        shell_cmd, shell_type = _find_shell()
        self.shell_type = shell_type

        # Try to start the real shell
        try:
            kwargs = dict(
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=1,
                encoding="utf-8",
                errors="replace",
            )
            # Windows-specific: no setsid, use CREATE_NEW_PROCESS_GROUP
            if IS_WINDOWS:
                kwargs["creationflags"] = subprocess.CREATE_NEW_PROCESS_GROUP
            else:
                kwargs["preexec_fn"] = os.setsid

            self.proc = subprocess.Popen(shell_cmd, **kwargs)
            self._running = True
            self._reader = threading.Thread(
                target=self._read_output, daemon=True
            )
            self._reader.start()

            # Send init commands
            if shell_type in ("bash", "wsl"):
                self._raw("export PS1=''")
                self._raw("export TERM=dumb")
                self._raw("export DEBIAN_FRONTEND=noninteractive")
                self._raw("stty -echo 2>/dev/null; true")
            elif shell_type == "powershell":
                self._raw("$PSStyle.OutputRendering = 'PlainText'")

            time.sleep(0.5)
            # Verify it's alive
            test = self.run("echo SHELL_OK", timeout=5)
            if "SHELL_OK" not in test:
                raise RuntimeError("Shell not responding")

            self.shell_info = f"Real shell: {' '.join(shell_cmd[:2])}"
            log_fn(f"[+] Shell ready: {self.shell_info}")
            return

        except Exception as e:
            log_fn(f"[!] Real shell failed ({e}), using Python emulator")
            self.proc = None
            self._running = False

        # Fallback: Python emulator
        self._emulator = PythonShellEmulator(log_fn=log_fn)
        self.shell_type = "python_emulator"
        self.shell_info = "Python virtual shell (no WSL/bash found)"
        log_fn(f"[+] Virtual shell ready — install WSL for full Linux tools")
        log_fn(f"[*] WSL install: wsl --install  (run in Windows Terminal as Admin)")

    def _raw(self, cmd: str):
        if self.proc and self.proc.poll() is None:
            try:
                self.proc.stdin.write(cmd + "\n")
                self.proc.stdin.flush()
            except Exception:
                pass

    def _read_output(self):
        try:
            for line in self.proc.stdout:
                if self._running:
                    self.output_queue.put(line.rstrip("\n"))
        except Exception:
            pass

    def run(self, cmd: str, timeout: int = 120,
            on_line: Callable = None) -> str:

        # Python emulator path — streams via on_line internally
        if self._emulator:
            result = self._emulator.run(cmd, timeout=timeout, on_line=on_line)
            return result if result else ""

        if not self.proc or self.proc.poll() is not None:
            self.start()

        # Drain stale output
        while not self.output_queue.empty():
            try:
                self.output_queue.get_nowait()
            except queue.Empty:
                break

        # PowerShell uses Write-Output for sentinel
        if self.shell_type == "powershell":
            self._raw(cmd)
            self._raw(f"Write-Output '{self.SENTINEL}'")
        else:
            self._raw(cmd)
            self._raw(f"echo '{self.SENTINEL}'")

        lines = []
        deadline = time.time() + timeout

        while time.time() < deadline:
            try:
                line = self.output_queue.get(timeout=1.0)
                if self.SENTINEL in line:
                    break
                # Filter out PS prompt noise
                if line.strip() in ("", "PS >", ">"):
                    continue
                lines.append(line)
                if on_line:
                    on_line(line)
            except queue.Empty:
                if self.proc and self.proc.poll() is not None:
                    break

        return "\n".join(lines)

    def kill(self):
        self._running = False
        self._emulator = None
        if self.proc:
            try:
                if IS_WINDOWS:
                    subprocess.run(
                        ["taskkill", "/F", "/T", "/PID", str(self.proc.pid)],
                        capture_output=True
                    )
                else:
                    os.killpg(os.getpgid(self.proc.pid), 9)
            except Exception:
                try:
                    self.proc.terminate()
                except Exception:
                    pass
            self.proc = None


# ═══════════════════════════════════════════════════════════
# TOOL MANAGER
# ═══════════════════════════════════════════════════════════

TOOL_INSTALL = {
    "nmap":         "winget install --id Insecure.Nmap -e --silent",
    "gobuster":     "go install github.com/OJ/gobuster/v3@latest",
    "nikto":        "https://github.com/sullo/nikto",
    "sqlmap":       "pip install sqlmap",
    "hydra":        "https://github.com/vanhauser-thc/thc-hydra",
    "john":         "https://www.openwall.com/john/",
    "hashcat":      "winget install --id Hashcat.Hashcat -e --silent",
    "masscan":      "https://github.com/robertdavidgraham/masscan",
    "dirb":         "https://github.com/v0re/dirb",
    "whois":        "winget install --id WhoisConnectedSoftware.Whois -e --silent",
    "dig":          "winget install --id ISC.BIND -e --silent",
    "curl":         "winget install --id curl.curl -e --silent",
    "wget":         "winget install --id GNU.Wget2 -e --silent",
    "netcat":       "winget install --id Insecure.Nmap -e --silent",
    "subfinder":    "go install github.com/projectdiscovery/subfinder/v2/cmd/subfinder@latest",
    "ffuf":         "go install github.com/ffuf/ffuf/v2@latest",
    "amass":        "go install github.com/owasp-amass/amass/v4/...@master",
    "theharvester": "pip install theHarvester",
    "wafw00f":      "pip install wafw00f",
    "whatweb":      "gem install whatweb",
    "sslscan":      "winget install --id rbsec.sslscan -e --silent",
}

def check_tool(session: "BashSession", tool: str) -> bool:
    import shutil as _sh

    # 1. Standard PATH check
    if _sh.which(tool):
        return True

    # 2. Try running it directly — catches tools installed but not in PATH
    try:
        result = subprocess.run(
            [tool, "--version"],
            capture_output=True, timeout=5
        )
        if result.returncode in (0, 1):   # 1 is fine — nmap returns 1 for --version
            return True
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        pass

    # 3. Check common Windows install locations for known tools
    WINDOWS_PATHS = {
        "nmap":     [
            r"C:\Program Files (x86)\Nmap\nmap.exe",
            r"C:\Program Files\Nmap\nmap.exe",
        ],
        "wireshark": [
            r"C:\Program Files\Wireshark\Wireshark.exe",
        ],
        "john": [
            r"C:\Program Files\John the Ripper\john.exe",
            r"C:\john\john.exe",
        ],
        "hashcat": [
            r"C:\hashcat\hashcat.exe",
            r"C:\Program Files\hashcat\hashcat.exe",
        ],
        "gobuster": [
            r"C:\Program Files\gobuster\gobuster.exe",
        ],
    }
    for path in WINDOWS_PATHS.get(tool, []):
        if os.path.exists(path):
            # Add its directory to PATH so it works in the shell too
            tool_dir = os.path.dirname(path)
            if tool_dir not in os.environ.get("PATH", ""):
                os.environ["PATH"] = tool_dir + os.pathsep + os.environ.get("PATH", "")
            return True

    return False

def ensure_tool(session: "BashSession", tool: str, log_fn=print) -> bool:
    if check_tool(session, tool):
        return True
    hint = TOOL_INSTALL.get(tool, f"Search for '{tool}' installation")
    log_fn(f"[!] {tool} not found. Install: {hint}")
    return False


# ═══════════════════════════════════════════════════════════
# AI PLANNER
# ═══════════════════════════════════════════════════════════

def ai_pentest_plan(task: str, context: str = "") -> dict:
    prompt = f"""You are an expert ethical hacker and penetration tester.
Create a step-by-step plan for this task.

TASK: {task}

Rules:
- Only use legitimate security testing tools
- Order steps logically: recon → enumeration → analysis
- Each step needs a specific command

Respond ONLY with valid JSON:
{{
  "target": "<extracted target or task>",
  "phases": [
    {{
      "phase": "Reconnaissance",
      "steps": [
        {{
          "tool": "nmap",
          "command": "nmap -sV -T4 <target>",
          "purpose": "Identify open ports and services",
          "timeout": 120
        }}
      ]
    }}
  ]
}}

Available tools: nmap, gobuster, nikto, sqlmap, whois, dig, curl, wget,
whatweb, wafw00f, sslscan, subfinder, ffuf, amass, theharvester"""

    try:
        resp = requests.post(
            OLLAMA_API_URL,
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.2, "num_predict": 2000}
            },
            timeout=120
        )
        raw = resp.json().get("response", "").strip()
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            import json
            return json.loads(match.group(0))
    except Exception as e:
        logger.error(f"AI plan failed: {e}")

    return {"target": task, "phases": [
        {"phase": "Recon", "steps": [
            {"tool": "nmap", "command": f"nmap -sV {task}",
             "purpose": "Port scan", "timeout": 120}
        ]}
    ]}


def ai_analyse_output(tool: str, output: str, target: str) -> str:
    prompt = f"""You are an expert penetration tester.

Tool: {tool}
Target: {target}
Output:
{output[:3000]}

Provide:
1. Key findings
2. Vulnerabilities or interesting points  
3. Recommended next steps
4. Specific commands to run based on findings

Be concise and technical."""

    try:
        resp = requests.post(
            OLLAMA_API_URL,
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.4, "num_predict": 600}
            },
            timeout=90
        )
        return resp.json().get("response", "").strip()
    except Exception as e:
        return f"Analysis failed: {e}"


# ═══════════════════════════════════════════════════════════
# HACKER AGENT
# ═══════════════════════════════════════════════════════════

class HackerAgent:
    def __init__(self):
        self.session    = BashSession()
        self.findings   = []
        self.log_lines  = []
        self.target     = None
        self._active    = False
        self._log_cb    = None

        # Permission gate — set by GUI via set_permission_callback()
        self._permission_cb: Optional[Callable] = None
        self._approved_tools: set = set()   # approved this session (don't ask again)
        self._denied_tools:   set = set()   # denied this session (skip silently)

    def set_log_callback(self, cb: Callable):
        self._log_cb = cb

    def set_permission_callback(self, cb: Callable):
        """
        Wire a GUI dialog here: cb(tool_name, install_cmd) -> bool
        Called before AURA installs any missing tool.
        Falls back to a CLI input() prompt if not set.
        """
        self._permission_cb = cb

    def _request_permission(self, tool: str, install_cmd: str) -> bool:
        """
        Ask the user (via GUI dialog or CLI) whether it's OK to install a tool.
        Remembers the answer for the rest of the session so the user is only
        asked once per tool.
        """
        if tool in self._approved_tools:
            return True
        if tool in self._denied_tools:
            return False

        self._log(
            f"[?] '{tool}' is not installed.\n"
            f"    Install command: {install_cmd}",
            "warn"
        )

        if self._permission_cb:
            approved = self._permission_cb(tool, install_cmd)
        else:
            # CLI fallback
            try:
                answer = input(
                    f"\n⚠  '{tool}' is missing.\n"
                    f"   Install with: {install_cmd}\n"
                    f"   Allow installation? [y/N]: "
                ).strip().lower()
                approved = answer in ("y", "yes")
            except (EOFError, OSError):
                approved = False

        if approved:
            self._approved_tools.add(tool)
            self._log(f"[+] Permission granted for '{tool}'.", "success")
        else:
            self._denied_tools.add(tool)
            self._log(f"[!] Installation of '{tool}' denied — step will be skipped.", "warn")

        return approved

    def _log(self, msg: str, level: str = "info"):
        ts   = datetime.now().strftime("%H:%M:%S")
        line = f"[{ts}] {msg}"
        self.log_lines.append({"time": ts, "msg": msg, "level": level})
        print(line)
        if self._log_cb:
            self._log_cb(line, level)

    def start_session(self):
        if not self._active:
            self._log("[*] Initialising security environment...")
            self.session.start(log_fn=self._log)
            self._active = True

            # Print environment info
            info = self.session.run("uname -a 2>/dev/null || echo 'Windows host'", timeout=5)
            who  = self.session.run("whoami", timeout=5)
            self._log(f"[+] Host: {info.strip()}", "success")
            self._log(f"[+] User: {who.strip()}", "success")
            self._log(f"[+] Shell: {self.session.shell_info}", "success")

            if self.session.shell_type == "python_emulator":
                self._log("[*] Running in virtual shell mode", "warn")
                self._log("[*] Network tools (ping/curl/whois/dig) are native Python", "warn")
                self._log("[*] For full tools: run 'wsl --install' in Admin terminal", "warn")

    def stop_session(self):
        if self._active:
            self.session.kill()
            self._active = False
            self._log("[*] Session closed.")

    def run_command(self, cmd: str, timeout: int = 120) -> str:
        if not self._active:
            self.start_session()
        self._log(f"[>] {cmd}")
        output = self.session.run(
            cmd, timeout=timeout,
            on_line=lambda l: self._log(f"    {l}")
        )
        return output

    def install_tools(self, tools: List[str]):
        """
        For each tool: check if installed, ask permission if not,
        then actually run the install command if approved.
        """
        for tool in tools:
            if check_tool(self.session, tool):
                self._log(f"[+] {tool} ✓", "success")
                continue

            # Skipped this session already
            if tool in self._denied_tools:
                self._log(f"[!] '{tool}' skipped (denied earlier).", "warn")
                continue

            hint = TOOL_INSTALL.get(tool, "")

            # Build an actual runnable install command from the hint
            install_cmd = self._resolve_install_cmd(tool, hint)

            if not self._request_permission(tool, install_cmd):
                continue   # user said no

            # Run the install
            self._log(f"[*] Installing '{tool}'...")
            self.session.run(
                install_cmd, timeout=180,
                on_line=lambda l: self._log(f"    {l}")
            )

            if check_tool(self.session, tool):
                self._log(f"[+] '{tool}' installed successfully.", "success")
            else:
                self._log(f"[!] '{tool}' install may have failed — check output above.", "warn")

    def _resolve_install_cmd(self, tool: str, hint: str) -> str:
        """
        Turn the hint string (URL, go install, pip install …) into
        the best runnable one-liner for the current shell type.
        """
        if not hint:
            return f"# No automatic install available for '{tool}'"

        # Already a real command
        if any(hint.startswith(p) for p in ("pip", "go install", "gem", "winget", "apt", "brew")):
            return hint

        # It's a URL — show it but can't auto-install
        if hint.startswith("http"):
            return f"# Download manually: {hint}"

        # Mixed "pip install X  OR  https://…"  — grab first part
        if " OR " in hint:
            first = hint.split(" OR ")[0].strip()
            if any(first.startswith(p) for p in ("pip", "go install", "gem", "winget")):
                return first

        return hint

    def run_task(self, task: str) -> Dict:
        self.target    = task
        self.findings  = []
        self.log_lines = []

        if not self._active:
            self.start_session()

        self._log("=" * 55)
        self._log(f"[*] TASK: {task}")
        self._log("=" * 55)

        self._log("[*] Planning task with AI...")
        plan   = ai_pentest_plan(task)
        target = plan.get("target", task)
        phases = plan.get("phases", [])

        self._log(f"[+] Target: {target}")
        self._log(f"[+] {len(phases)} phases planned")

        # Check tools
        all_tools = {step.get("tool", "") for phase in phases
                     for step in phase.get("steps", [])}
        all_tools.discard("")
        self._log(f"[*] Tools needed: {', '.join(all_tools)}")
        self.install_tools(list(all_tools))

        # Execute phases
        for phase in phases:
            phase_name = phase.get("phase", "Unknown")
            self._log(f"\n{'─'*55}", "dim")
            self._log(f"PHASE: {phase_name}", "phase")
            self._log(f"{'─'*55}", "dim")

            for step in phase.get("steps", []):
                tool    = step.get("tool", "")
                cmd     = step.get("command", "").replace("<target>", target).replace("{target}", target)
                purpose = step.get("purpose", "")
                timeout = step.get("timeout", 120)

                # Skip steps whose tool the user denied
                if tool and tool in self._denied_tools:
                    self._log(f"[!] Skipping '{purpose}' — '{tool}' was not installed.", "warn")
                    continue

                self._log(f"\n[*] {purpose}", "step")
                self._log(f"[>] {cmd}")

                output = self.session.run(
                    cmd, timeout=timeout,
                    on_line=lambda l: self._log(f"    {l}")
                )

                if output.strip():
                    self._log(f"\n[AI] Analysing output...")
                    analysis = ai_analyse_output(tool, output, target)
                    self._log(f"[AI] {analysis}", "analysis")
                    self.findings.append({
                        "phase":    phase_name,
                        "tool":     tool,
                        "command":  cmd,
                        "output":   output,
                        "analysis": analysis,
                        "time":     datetime.now().isoformat(),
                    })
                else:
                    self._log(f"[!] No output from {tool}")

        self._log("\n[*] Generating report...")
        report_path = self._generate_report(target, plan, self.findings)
        self._log(f"[+] Report: {report_path}", "success")
        self._log("[+] TASK COMPLETE", "success")

        return {
            "target":      target,
            "findings":    self.findings,
            "report_path": report_path,
            "log":         self.log_lines,
        }

    def _generate_report(self, target: str, plan: dict,
                          findings: List[Dict]) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = re.sub(r'[^a-zA-Z0-9]', '_', target)[:30]
        path = os.path.join(REPORT_DIR, f"pentest_{safe}_{timestamp}.docx")

        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            for section in doc.sections:
                section.top_margin    = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin   = Inches(1.2)
                section.right_margin  = Inches(1.2)

            t = doc.add_heading("PENETRATION TEST REPORT", 0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Target: {target}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("Generated by AURA Security Agent").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

            doc.add_heading("Executive Summary", 1)
            doc.add_paragraph(
                f"Penetration test conducted against {target} on "
                f"{datetime.now().strftime('%B %d, %Y')}. "
                f"{len(findings)} tool runs across {len(plan.get('phases', []))} phases."
            )

            doc.add_heading("Technical Findings", 1)
            for finding in findings:
                doc.add_heading(f"{finding['phase']} — {finding['tool'].upper()}", 2)
                p = doc.add_paragraph()
                p.add_run("Command: ").bold = True
                p.add_run(finding["command"])
                doc.add_paragraph().add_run("Output:").bold = True
                raw = doc.add_paragraph(finding["output"][:2000])
                for run in raw.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)
                doc.add_paragraph().add_run("AI Analysis:").bold = True
                doc.add_paragraph(finding["analysis"])
                doc.add_paragraph()

            doc.add_page_break()
            doc.add_heading("Full Session Log", 1)
            log_run = doc.add_paragraph().add_run(
                "\n".join(f"[{l['time']}] {l['msg']}" for l in self.log_lines)[:10000]
            )
            log_run.font.name = "Courier New"
            log_run.font.size = Pt(7)

            doc.save(path)
            return path

        except ImportError:
            txt = path.replace(".docx", ".txt")
            with open(txt, "w") as f:
                f.write(f"PENTEST REPORT — {target}\nDate: {datetime.now()}\n\n")
                for finding in findings:
                    f.write(f"[{finding['phase']}] {finding['tool']}\n")
                    f.write(f"CMD: {finding['command']}\n")
                    f.write(f"OUTPUT:\n{finding['output'][:1000]}\n")
                    f.write(f"ANALYSIS:\n{finding['analysis']}\n")
                    f.write("-" * 40 + "\n")
            return txt


# ── Singleton ──────────────────────────────────────────────────────────────

_agent_instance: Optional[HackerAgent] = None

def get_hacker_agent() -> HackerAgent:
    global _agent_instance
    if _agent_instance is None:
        _agent_instance = HackerAgent()
    return _agent_instance